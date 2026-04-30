import docx
import os

def add_chapter_4():
    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch3_Persistence_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_page_break()

    # --- Chapter 4 ---
    doc.add_heading('Chapter 4\nImplementation, Testing, and Maintenance', level=1)
    
    doc.add_heading('4.1 Introduction to Languages, IDEs, Tools and Technologies', level=2)
    doc.add_paragraph("Unlike basic static websites, AtmoSense AI is engineered to function as a dynamic, highly responsive environmental intelligence platform. The technology stack was deliberately selected to maximize asynchronous execution speed, seamless third-party API communication, and interactive UI rendering.")
    
    doc.add_heading('Core Programming Language and Interface', level=3)
    doc.add_paragraph("JavaScript (ES6+) and Node.js: The entirety of the AtmoSense AI ecosystem, from its backend services layer to its frontend components, is written in JavaScript. Node.js was chosen due to its unparalleled asynchronous event-driven architecture, enabling the server to handle parallel API requests to external weather providers without blocking the main thread.")
    doc.add_paragraph("React.js and Tailwind CSS: Instead of a rigid, multi-page HTML architecture, AtmoSense utilizes a lightweight, component-based React frontend. Tailwind CSS is used for utility-first styling, ensuring zero custom CSS overhead while delivering a modern, glassmorphism-styled UI that feels premium and responsive.")

    doc.add_heading('Environmental Data Integration (The API Gateway)', level=3)
    doc.add_paragraph("OpenWeatherMap API: Serving as the primary meteorological engine, this API is utilized to fetch real-time temperature, humidity, and wind metrics. It guarantees low-latency global coverage for accurate heat index calculations.")
    doc.add_paragraph("WAQI (World Air Quality Index) API: Serving as the secondary environmental engine. The WAQI API is utilized to retrieve dense, granular pollutant metrics (PM2.5, PM10) necessary for calculating lung recovery times.")
    doc.add_paragraph("Mock Data Fallback System: Integrated as a tertiary fallback provider, ensuring the frontend dashboard maintains 100% uptime even during catastrophic third-party API rate limits or network outages.")

    doc.add_heading('Database & Backend Modules', level=3)
    doc.add_paragraph("MongoDB and Mongoose: A scalable NoSQL database utilized extensively by the backend module to persist structured BSON documents (PredictionLogs). Mongoose enforces strict schema validation before any data is written to the Atlas cloud cluster.")
    doc.add_paragraph("Axios & Express.js: Axios is used to securely spawn parallel HTTP GET requests to external providers, while Express.js orchestrates the REST API endpoints and middleware logic for the React client.")

    doc.add_heading('Development Tools and Environment', level=3)
    doc.add_paragraph("Visual Studio Code (VS Code): The primary Integrated Development Environment used for architecting the MERN stack application, managing the Git repository, and debugging the Node.js event loop.")
    doc.add_paragraph("Vite: The modern frontend build tool used to serve the React application locally with ultra-fast Hot Module Replacement (HMR).")

    doc.add_heading('4.2 Testing Techniques and Test Plans', level=2)
    doc.add_paragraph("Due to the dynamic nature of AtmoSense AI—which relies on constantly shifting external data—rigorous testing protocols were implemented to prevent catastrophic frontend crashes.")

    doc.add_heading('Unit Testing & Network Resilience', level=3)
    doc.add_paragraph("The Backend API Gateway was subjected to extensive isolated testing. Network requests were intentionally malformed or repeatedly spammed to trigger HTTP 429 Rate Limit and HTTP 403 Access Denied exceptions. The unit tests verified that the Express system successfully caught these exceptions and seamlessly pivoted to the internal Mock Data Generator without crashing the Node.js process.")

    doc.add_heading('Heuristic Algorithms Validation', level=3)
    doc.add_paragraph("The Prediction Engine module employs complex mathematical equations (like the Rothfusz regression). To test this, the system passes extreme boundary values (e.g., 50°C temperature and 100% humidity). The validation logic ensures that the algorithms do not output NaN (Not a Number) or infinite values, preventing the corruption of the JSON response payload.")

    doc.add_heading('Integration Testing (Database & Frontend)', level=3)
    doc.add_paragraph("Integration tests were conducted to verify the alignment between the React state manager and the MongoDB database. Tests confirmed the ability to successfully submit a user profile, calculate heuristics, and persist the record to MongoDB Atlas.")

    doc.add_heading('Sample Test Cases', level=3)
    doc.add_paragraph("Table 4.1 – Test Cases for Core Operations:")
    doc.add_paragraph("TC-01 (API Fallback): Input – OpenWeatherMap API rate limit deliberately exceeded. Expected Output – Services Layer catches the 429 error, logs a warning, and routes the request to the synthetic data generator. Result – Pass.")
    doc.add_paragraph("TC-02 (Heuristic Execution): Input – User submits 'Running' activity with 35°C ambient temperature. Expected Output – System calculates elevated heat index, prescribes 950ml/hr hydration, and returns JSON. Result – Pass.")
    doc.add_paragraph("TC-03 (Database Persistence): Input – Valid prediction payload. Expected Output – Mongoose schema validates the data types and successfully commits a new document to the Atlas cluster. Result – Pass.")
    doc.add_paragraph("TC-04 (UI Rendering): Input – System returns a \"DANGER\" safety level. Expected Output – React components conditionally render a red background and trigger Framer Motion alert animations. Result – Pass.")

    doc.add_heading('4.3 Installation Instructions', level=2)
    doc.add_paragraph("Step 1: Clone the AtmoSense AI project repository securely from GitHub to the local machine.")
    doc.add_paragraph("Step 2: Ensure Node.js (Version 18 or higher) and NPM are installed natively on the host operating system, and verify that they are properly added to the system PATH environment variables.")
    doc.add_paragraph("Step 3: Navigate to the atmosense-backend directory via terminal and execute 'npm install' to install Express, Mongoose, Axios, and dotenv dependencies.")
    doc.add_paragraph("Step 4: System Configuration – Create a '.env' file in the backend root directory. Securely inject your required credentials: OPENWEATHER_API_KEY, WAQI_API_KEY, and MONGODB_URL. Ensure this file is listed in .gitignore.")
    doc.add_paragraph("Step 5: Navigate to the atmosense-frontend directory and execute 'npm install' to resolve React, Vite, Tailwind CSS, and Framer Motion dependencies.")

    doc.add_heading('4.4 End User Instructions', level=2)
    doc.add_paragraph("Follow these steps to successfully operate the AtmoSense AI application:")
    doc.add_paragraph("Step 1: Open a terminal application with standard user privileges.")
    doc.add_paragraph("Step 2: Navigate to the backend directory and launch the Express server by executing the command: 'npm start' (or 'node index.js').")
    doc.add_paragraph("Step 3: Open a second terminal window, navigate to the frontend directory, and start the development server using: 'npm run dev'.")
    doc.add_paragraph("Step 4: Upon execution, open your modern web browser (e.g., Chrome, Edge) and navigate to http://localhost:5173.")
    doc.add_paragraph("Step 5: The browser will render the interactive, glassmorphism-styled AtmoSense AI dashboard.")
    doc.add_paragraph("Step 6: Enter your target city, select your planned physical activity, and define your health sensitivity using the provided dropdown menus.")
    doc.add_paragraph("Step 7: Click the 'Predict Health Impact' button and allow the system to autonomously gather data and compute physiological heuristics.")
    doc.add_paragraph("Step 8: Review the final rendered output. The system will display visually color-coded safety metrics, exact hydration requirements, and pollutant breakdowns directly on your screen.")

    output_path = r'Report\AtmoSense_Ch1_to_Ch4_Final.docx'
    doc.save(output_path)
    print("Chapter 4 successfully appended to the document.")

if __name__ == '__main__':
    add_chapter_4()
