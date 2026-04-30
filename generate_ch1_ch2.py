import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def create_report():
    doc = docx.Document()
    
    # --- Title Page ---
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('AtmoSense AI\n(Smart Environmental Health Intelligence)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(24)
        
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('A Minor Project Report\nSubmitted in partial fulfillment of requirement of the\nDegree of\nBACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING\nBY')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Piyush Solanki (EN23CS304044)\nMayank Sahu (EN23CS304037)\nPiyush Pal (EN23CS304044)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Under the Guidance of\nProf. Vishal Sharma\nProf. Suyog Munshi')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Department of Computer Science & Engineering\nFaculty of Engineering\nMEDICAPS UNIVERSITY, INDORE- 453331\n\nAPRIL-2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Chapter 1 ---
    doc.add_heading('Chapter 1\nIntroduction', level=1)
    
    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph("Environmental tracking and personal health management play a critical role in modern daily routines. Every day, individuals and athletes across various regions execute complex outdoor activities that require understanding changing weather patterns, air quality indices, and their direct physiological impacts. Managing personal health in extreme conditions, predicting hydration needs, and assessing lung recovery times manually is both time-consuming and repetitive, particularly when dealing with rapidly fluctuating environmental variables.")
    doc.add_paragraph("AtmoSense AI is an AI-powered environmental health intelligence ecosystem designed to address these challenges. The system intelligently analyzes real-time meteorological data, air quality metrics, and user-specific physiological intent to automatically generate context-aware health heuristics using advanced predictive algorithms. By integrating environmental data directly into a unified interactive dashboard, AtmoSense AI eliminates the need to switch between multiple weather and health apps, thereby improving decision-making efficiency and reducing health risks.")
    doc.add_paragraph("The system is built using a modern, robust technology stack: Node.js and Express for core backend execution and API orchestration, a modular React.js interface with Tailwind CSS for frontend interaction, and a highly resilient AI prediction engine for intelligent impact calculation. AtmoSense AI is deployed as a full-stack web application, enabling users to seamlessly track AQI, predict heat stress, and receive customized safety recommendations directly within their browser in a secure manner.")
    
    doc.add_heading('1.2 Literature Review', level=2)
    doc.add_paragraph("Recent advancements in Data Analytics and Health Informatics have enabled the development of automated systems capable of not just displaying weather data, but reasoning through complex physiological logic. Models such as the Rothfusz regression for Heat Index and generalized lung recovery heuristics have demonstrated remarkable capabilities in understanding human thermoregulation and respiratory stress.")
    doc.add_paragraph("Several weather tracking tools have been developed to assist users in managing outdoor activities. Applications like AccuWeather and AQICN introduced predictive forecasting and air quality tracking. However, these features offer limited personalization; they exist in isolated data silos and suffer from the \"interpretation gap.\" They generate numerical data, but still require the user to manually calculate hydration needs, assess safety limits, and adjust workout intensity.")
    doc.add_paragraph("Studies in the domain of environmental health highlight the growing need for context-sensitive, autonomous health predictions rather than passive data dashboards. The reliance on standalone weather APIs also presents significant vulnerabilities, as users frequently encounter downtime or lack localized data during intensive outdoor planning.")
    doc.add_paragraph("AtmoSense AI builds upon this body of work by offering a truly integrated health platform that combines advanced capabilities through a proprietary AI Prediction Engine. This approach overcomes the key limitations of existing tools—specifically, their lack of personalized health heuristics, susceptibility to API failures, and poor synthesis of combined weather and pollution data.")

    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph("To develop an AI-powered environmental intelligence platform that helps users efficiently plan outdoor activities with minimal health risks. To significantly reduce the cognitive effort required to interpret complex weather and AQI data by automating the calculation process using advanced predictive algorithms. To integrate the health prediction feature directly into a unified dashboard, ensuring seamless usage without switching between disparate health and weather applications. To design a highly resilient backend architecture that automatically falls back to intelligent mock data to prevent downtime caused by external API limits. To build a scalable and maintainable modular MERN stack application that can securely interface with external REST APIs like OpenWeatherMap and WAQI. To allow users to customize their physiological profiles (e.g., Asthma sensitivity) and automate daily health planning autonomously.")

    doc.add_heading('1.4 Significance', level=2)
    doc.add_paragraph("The significance of AtmoSense AI lies in its ability to transform the way individuals interact with their immediate environment. Outdoor activities and personal health management remain foundational to human well-being, and the burden of manually interpreting AQI risks and heat indices has become a major source of health hazards.")
    doc.add_paragraph("AtmoSense AI addresses this challenge by providing an intelligent system that acts as a proactive health partner. This reduces the cognitive effort required to cross-reference multiple data sources, minimizes exposure to dangerous pollutants, and helps maintain a safe physiological state during complex physical activities.")
    doc.add_paragraph("Furthermore, by operating as a seamless web platform, AtmoSense AI integrates non-intrusively into the user's existing daily routine, making it accessible securely from any device. The system's fallback capability makes it uniquely reliable for uninterrupted health planning.")

    doc.add_heading('1.5 Research Design / Methodology', level=2)
    doc.add_paragraph("Requirement Analysis: Identifying the functional and non-functional requirements of the system through an analysis of user friction points (data silos, lack of personalization, manual health calculations).")
    doc.add_paragraph("System Design: Designing the overall modular MERN architecture, data flow, API orchestration mechanisms, fallback logic, and glassmorphism-styled UI/UX of the React frontend.")
    doc.add_paragraph("Implementation: Developing the core logic using Node.js and Express, integrating external REST APIs for environmental data gathering, utilizing specialized heuristics (e.g., Rothfusz regression) for health impact calculation, and connecting MongoDB for prediction logging.")
    doc.add_paragraph("Testing: Conducting unit testing of the heuristic prediction engine, integration testing of the OpenWeatherMap and WAQI APIs, and simulated network failure testing to validate the mock data fallback logic.")
    doc.add_paragraph("Deployment: Packaging the project for local execution, ensuring .env security isolation, and running the application within the user's standard browser environment. The system follows a decoupled architecture where the React frontend communicates dynamically with the Node backend, which securely interfaces with remote data APIs to generate health insights.")

    doc.add_heading('1.6 Source of Data', level=2)
    doc.add_paragraph("The primary source of data for AtmoSense AI is the real-time environmental APIs accessed dynamically. The backend fetches meteorological data (temperature, humidity, wind) and air quality metrics (PM2.5, PM10, NO2) to generate context.")
    doc.add_paragraph("For the prediction module, internal heuristics are gathered dynamically via established physiological models (e.g., National Weather Service Heat Index equations). No external personal datasets are used for model training, as the system relies entirely on zero-shot inference from algorithmic rules via the prediction engine.")
    doc.add_paragraph("Official documentation and technical references used during development include the React Official Documentation, Node.js API Guidelines, Mongoose SDK Documentation, Tailwind CSS specs, and the REST API references for OpenWeatherMap and WAQI.")

    doc.add_heading('1.7 Chapter Scheme', level=2)
    doc.add_paragraph("The report is organized into the following chapters:\n• Chapter 1 – Introduction: Presents the background, literature review, objectives, significance, methodology, and chapter organization of the project.\n• Chapter 2 – Requirements Specification: Details the functional and non-functional requirements, user characteristics, hardware specifications, constraints, and dependencies of the system.\n• Chapter 3 – Design: Describes the system architecture, data flow diagrams, activity diagrams, class diagrams, ER diagrams, sequence diagrams, and database/local storage design.\n• Chapter 4 – Implementation, Testing, and Maintenance: Covers the technologies used, heuristic predictive strategies, fallback testing, installation procedures, and end-user instructions.\n• Chapter 5 – Results and Discussions: Presents the React user interface, module descriptions, and system snapshots with brief explanations of the core features.\n• Chapter 6 – Summary and Conclusions: Summarizes the work done and presents conclusions drawn from the AtmoSense AI project.\n• Chapter 7 – Future Scope: Outlines the potential enhancements, mobile support, and future directions for the AtmoSense AI system.")
    doc.add_page_break()

    # --- Chapter 2 ---
    doc.add_heading('Chapter 2\nRequirements Specification', level=1)
    
    doc.add_heading('2.1 User Characteristics', level=2)
    doc.add_paragraph("Health-Conscious Individuals and Athletes: Professionals and hobbyists who manage rigorous outdoor routines, execute physical workouts, and require precision regarding environmental limits. They require the system to instantly analyze complex atmospheric data and output seamless physiological parameters like precise hydration requirements and body temperature elevations.")
    doc.add_paragraph("Medical Patients and Sensitive Demographics: Users responsible for maintaining their respiratory health, such as asthmatics. They rely on the system to autonomously monitor pollutant spikes (PM2.5, PM10) and aggressively calculate lung recovery times, ensuring they avoid hazardous outdoor exposure.")
    doc.add_paragraph("General Public and Commuters: Users who require robust daily productivity and safety automation. They depend on the system to parse current meteorological conditions to generate structured safety recommendations, assisting with routine schedule planning around peak heat indices.")
    doc.add_paragraph("Tech Enthusiasts and Developers: Individuals seeking to interact with a modern MERN application, explore its glassmorphism UI, and appreciate the underlying API fallback intelligence. They utilize the platform's robust backend to observe how asynchronous JavaScript orchestration efficiently resolves parallel data streams.")

    doc.add_heading('2.2 Functional Requirements', level=2)
    doc.add_paragraph("REQ-1: The system shall provide a unified, highly modular, React-based dashboard interface via the frontend module, allowing users to select from distinct operational features like prediction and air quality tracking without leaving the single-page application.")
    doc.add_paragraph("REQ-2: The system shall implement a proprietary 'Mock Data Fallback' mechanism. If the primary external APIs (e.g., OpenWeatherMap or WAQI) encounter an HTTP 429 Rate Limit or HTTP 403 Access Denied error, the system must automatically and transparently failover to an internal intelligent data generator.")
    doc.add_paragraph("REQ-3: The system shall possess deep HTTP integration, utilizing Node.js and Axios modules to safely execute parallel GET requests, normalize incoming JSON payloads, and manage asynchronous data streams efficiently.")
    doc.add_paragraph("REQ-4: The system (Prediction Engine module) shall read the normalized weather and AQI structures, and apply heuristic, algorithmic equations (such as the Rothfusz regression) directly to the data to resolve runtime health impact calculations.")
    doc.add_paragraph("REQ-5: The system (Backend module) shall dynamically detect incoming POST requests containing user profiles, initialize a Mongoose schema, push to MongoDB Atlas via secured connection strings, and orchestrate REST API responses for the frontend.")
    doc.add_paragraph("REQ-6: The system shall extract specific environmental variables (Temperature, PM2.5, Humidity) dynamically, feeding the extracted values to the internal algorithms for immediate processing of hydration and lung recovery times.")
    doc.add_paragraph("REQ-7: The system shall aggregate data by parsing REST APIs, generating a highly personalized, visual safety recommendation dashboard tailored specifically to the user's defined physical activity and health sensitivity.")
    doc.add_paragraph("REQ-8: The system shall process complex datasets and utilize UI libraries like Recharts to generate structured bar charts, which are subsequently rendered flawlessly on the Air Quality dashboard page.")
    doc.add_paragraph("REQ-9: The system shall persist state strictly via MongoDB, while storing encrypted API credentials within a .env file locally, ensuring zero exposure of sensitive keys to the client-side browser.")

    doc.add_heading('2.3 Dependencies', level=2)
    doc.add_paragraph("Core Language: JavaScript (ES6+) / Node.js 18 or higher. The entire backend core and React frontend strictly rely on the Node runtime environment and V8 JavaScript engine.")
    doc.add_paragraph("External Application Programming Interfaces (APIs): The system's logical reasoning depends heavily on external API endpoints, specifically the OpenWeatherMap Current Weather API and the World Air Quality Index (WAQI) JSON API.")
    doc.add_paragraph("Backend Modules: The system extensively uses built-in and third-party Node.js libraries such as Express.js for routing, Cors for cross-origin policy management, and dotenv to natively bind to local environment variables securely.")
    doc.add_paragraph("Database Libraries: Mongoose ODM is strictly required for modeling application data and interacting with the MongoDB cluster. The 'mongodb-memory-server' is also utilized as a dependency for local, in-memory database testing and fallback operations.")
    doc.add_paragraph("Frontend UI Libraries: React.js and Vite are necessary for DOM manipulation and bundling. Tailwind CSS is required for utility-first styling. Framer Motion is essential for fluid component animations, and Lucide React is required for scalable vector icon rendering.")
    doc.add_paragraph("Data Processing Libraries: Axios is necessary for executing HTTP requests and manipulating JSON payloads efficiently from the backend server to the external providers.")

    doc.add_heading('2.4 Performance Requirements', level=2)
    doc.add_paragraph("• The API orchestration layer must detect external API rate limits or failures, intercept the exception, and successfully route the request to the Mock Data generator within 500 milliseconds to prevent user-facing latency.")
    doc.add_paragraph("• System data queries to OpenWeatherMap and WAQI must resolve concurrently and render to the React frontend within 1.5 seconds under standard network conditions.")
    doc.add_paragraph("• Algorithmic heuristic calculations (Heat Index, Hydration, Lung Recovery) must analyze, compute, and inject results into the JSON response payload within 50 to 100 milliseconds.")
    doc.add_paragraph("• The daily dashboard generation, which involves rendering complex Recharts SVG graphs and Framer Motion animations, must maintain a rendering frame rate of 60 FPS (Frames Per Second) on the client side.")

    doc.add_heading('2.5 Hardware Requirements', level=2)
    doc.add_paragraph("Processor: Multi-core CPU (Intel Core i3, AMD Ryzen 3, or Apple Silicon equivalent) to handle concurrent Node.js event-loop operations and local Vite server compiling.")
    doc.add_paragraph("RAM: Minimum 4 GB (8 GB highly recommended to allow AtmoSense AI to monitor the system while heavy IDEs and background development servers are running).")
    doc.add_paragraph("Storage: Minimum 1 GB of free disk space for Node_modules, installed npm dependencies, and MongoDB local caching.")
    doc.add_paragraph("Network: Stable broadband internet connection (minimum 5 Mbps) to ensure low-latency communication with the various external APIs and MongoDB Atlas Cloud cluster.")
    doc.add_paragraph("Operating System: Windows, macOS, or Linux is required for backend hosting. On the client side, any modern web browser (Chrome, Edge, Safari, Firefox) is strictly required for full UI functionality.")

    doc.add_heading('2.6 Constraints and Assumptions', level=2)
    doc.add_paragraph("Constraints:")
    doc.add_paragraph("• The backend logic must be exclusively implemented in Node.js/Express, adhering to the established decoupled Client-Server architectural paradigm.")
    doc.add_paragraph("• The user interface must remain entirely web-based; no native desktop binaries or terminal-only execution is permitted for end-user operations.")
    doc.add_paragraph("• Sensitive API keys must never be hardcoded or pushed to version control; they must exclusively reside in a git-ignored local .env file on the server.")
    doc.add_paragraph("• The MongoDB implementation strictly constrains the user to either provide a valid Atlas URI or rely on the temporary local memory server.")
    doc.add_paragraph("Assumptions:")
    doc.add_paragraph("• Users possess a modern web browser capable of rendering CSS Grid, Flexbox, and complex JavaScript animations without lag.")
    doc.add_paragraph("• The host machine (for local deployment) has Node.js and NPM added to the system PATH environment variable, allowing the execution of 'npm run dev' and 'node index.js' from any directory.")
    doc.add_paragraph("• For live data, users have successfully generated and configured active API keys for OpenWeatherMap and WAQI inside the backend environment configuration.")

    doc.save(r'Report\AtmoSense_Ch1_Ch2_Detailed_v2.docx')

if __name__ == '__main__':
    create_report()
