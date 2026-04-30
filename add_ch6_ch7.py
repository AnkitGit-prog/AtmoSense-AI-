import docx
import os

def append_conclusions():
    doc_path = r'Report\AtmoSense_Ch1_to_Ch5_Complete_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_page_break()

    # --- Chapter 6 ---
    doc.add_heading('Chapter 6\nSummary and Conclusions', level=1)
    
    doc.add_heading('6.1 Summary', level=2)
    doc.add_paragraph("AtmoSense AI was developed to bridge the gap between raw environmental data and actionable personal health heuristics. By utilizing the modern MERN stack (MongoDB, Express, React, Node.js), the application provides a highly responsive, asynchronous, and visually intuitive platform that allows individuals to monitor their surroundings in real time. The integration of OpenWeatherMap and WAQI REST APIs ensures that the meteorological and pollutant data is consistently accurate and geographically relevant. Furthermore, the inclusion of a proprietary Mock Data Fallback architecture guarantees platform resilience, masking any backend API rate limits or network outages from the end user.")
    doc.add_paragraph("The frontend, styled with utility-first Tailwind CSS and animated via Framer Motion, succeeds in transforming complex atmospheric data into easily readable safety badges, precise hydration metrics, and interactive Recharts. This significantly reduces the cognitive load on the user, enabling rapid, informed decisions regarding outdoor activities.")

    doc.add_heading('6.2 Conclusions', level=2)
    doc.add_paragraph("The successful implementation of AtmoSense AI demonstrates the immense potential of combining robust API orchestration with targeted physiological algorithms. The system effectively processes complex, overlapping environmental variables (like PM2.5 concentrations and Heat Index) and outputs precise, user-centric safety recommendations.")
    doc.add_paragraph("By persisting all prediction logs to MongoDB Atlas, the platform not only acts as a real-time monitor but also establishes a secure foundation for historical health auditing. The project successfully fulfills all initial functional and non-functional requirements, delivering a robust, secure, and highly available full-stack application capable of protecting users from hazardous environmental exposure.")

    doc.add_page_break()

    # --- Chapter 7 ---
    doc.add_heading('Chapter 7\nFuture Scope', level=1)
    
    doc.add_heading('7.1 Machine Learning Integration', level=2)
    doc.add_paragraph("While the current AI Prediction Engine relies on established deterministic mathematical heuristics (e.g., the Rothfusz regression equation for Heat Index), future iterations of AtmoSense AI will integrate true Machine Learning (ML) models. By training algorithms such as Random Forest or Long Short-Term Memory (LSTM) networks on the historical data persisted in MongoDB, the system will be able to predict localized pollution spikes and temperature anomalies before they occur, shifting the platform from a reactive monitor to a proactive forecaster.")

    doc.add_heading('7.2 Wearable Device Synchronization', level=2)
    doc.add_paragraph("To further personalize the health impact predictions, future development will focus on integrating Web Bluetooth APIs and OAuth 2.0 protocols to sync real-time biometric data directly from wearable devices (e.g., Apple Watch, Fitbit, Garmin). By ingesting live heart rate and blood oxygen levels into the Express backend, the Prediction Engine will be able to dynamically adjust its hydration and lung recovery heuristics based on the user's immediate physiological stress.")

    doc.add_heading('7.3 Mobile Application Deployment', level=2)
    doc.add_paragraph("Currently implemented as a responsive React.js web application, the next logical step for the ecosystem is a port to React Native. This architectural shift will enable native application deployment on iOS and Android devices. Native integration will provide critical functionality such as background geolocation tracking and real-time push notifications, instantly alerting users to severe AQI warnings or extreme heat advisories directly on their lock screens.")

    output_path = r'Report\AtmoSense_Final_Project_Report.docx'
    doc.save(output_path)
    print("Chapters 6 and 7 successfully appended. Full report generated.")

if __name__ == '__main__':
    append_conclusions()
