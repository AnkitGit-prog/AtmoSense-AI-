import docx
import os

def append_activity_diagram():
    doc_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_With_Images.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.3.2 Activity Diagram', level=3)
    
    doc.add_paragraph("The Activity Diagram illustrates the sequential, state-driven flow of tasks originating from the user's initial interaction with the React Frontend to the final rendering of health heuristics on the client browser. Unlike traditional weather applications that follow a simple fetch-and-display loop, AtmoSense AI operates as a continuous, multi-state algorithmic pipeline.")
    doc.add_paragraph("The activity sequence is systematically broken down into the following operational phases:")
    
    doc.add_paragraph("1. React Initialization & Component Activation: The sequence begins when the user launches the AtmoSense AI web application. The system renders an interactive, glassmorphism-styled dashboard. Upon the user selecting a specific operation (e.g., AI Health Prediction or AQI Monitor), the React Router instantiates the corresponding page component.")
    doc.add_paragraph("2. Environmental & Physiological Context Gathering: Before consulting the AI algorithms, the system autonomously gathers local and user-defined state data to build a highly accurate execution context. Depending on the active module, this activity involves utilizing React State to capture physical activity types (Running/Walking), health sensitivities (Asthma), and extracting real-time geolocation or user-inputted city parameters.")
    doc.add_paragraph("3. Payload Construction & API Forwarding: The gathered physiological context is sanitized and injected into a structured JSON payload. The Axios client subsequently forwards this payload to the Express.js Backend Gateway, which initiates secure parallel REST GET requests to the primary environmental data providers (OpenWeatherMap and WAQI).")
    doc.add_paragraph("4. Network Resilience & Decision Node: A critical decision node occurs during API communication within the backend services layer. If the primary models return an HTTP 200 OK, the flow proceeds normally, aggregating real-time PM2.5 and temperature data. However, if the API throws an HTTP 429 Rate Limit or timeout exception, the activity flow branches into a fallback loop, automatically routing the data request to an internal Mock Data Generator to ensure uninterrupted user experience.")
    doc.add_paragraph("5. Heuristic Validation & Computation: Once the raw environmental data is successfully fetched or mocked, the system enters a computational state. Pure JavaScript functions strictly evaluate the Rothfusz regression equation for Heat Index, hydration scaling formulas, and lung recovery models. The algorithms verify that the generated heuristic bounds (e.g., Heat Index > 32°C) are logically valid and will not corrupt the final payload.")
    doc.add_paragraph("6. Final Rendering & Database Integration: Following successful computation, the final activity involves applying the changes to the user's view. The backend utilizes asynchronous Mongoose operations to persist the PredictionLog securely to MongoDB Atlas, while concurrently resolving the HTTP response back to the React frontend, where Framer Motion and Recharts dynamically render the color-coded safety badges.")

    output_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Complete.docx'
    doc.save(output_path)
    print("Activity Diagram section successfully appended.")

if __name__ == '__main__':
    append_activity_diagram()
