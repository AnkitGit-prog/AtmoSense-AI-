import docx
from docx.shared import Inches
import os

def insert_images_to_doc():
    doc_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Detailed.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('Visual Workflow Diagrams', level=2)
    doc.add_paragraph('Below are the generated visual workflow diagrams for the system architecture and data flow:')

    # Insert images that were successfully downloaded
    image_files = [f for f in os.listdir('Report') if f.endswith('.png') and f.startswith('Workflow_Diagram_')]
    
    for img in sorted(image_files):
        img_path = os.path.join('Report', img)
        doc.add_paragraph(f'Diagram: {img.replace(".png", "")}')
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            doc.add_paragraph('\n')
        except Exception as e:
            print(f"Could not add {img}: {e}")

    output_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_With_Images.docx'
    doc.save(output_path)
    print("Images successfully inserted into the Word document.")

if __name__ == '__main__':
    insert_images_to_doc()
