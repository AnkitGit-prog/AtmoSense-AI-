import docx
import os

def add_persistence():
    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch3_SequenceDiagram_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.4 NoSQL Data Persistence Design', level=2)
    
    doc.add_paragraph("Unlike traditional web applications that rely on heavy, rigid relational database management systems (RDBMS) like MySQL, AtmoSense AI prioritizes execution speed, schema flexibility, and cloud portability. Consequently, the system avoids traditional SQL databases entirely. Instead, the persistence layer is implemented using a scalable, NoSQL-style architecture utilizing strictly formatted BSON documents via MongoDB Atlas, combined with encrypted .env configurations for security.")
    
    doc.add_heading('3.4.1 Logical Storage Design', level=3)
    doc.add_paragraph("The logical design of AtmoSense AI consists of isolated document collections rather than relational tables. This design reflects the operational requirements of an intelligent environmental tracking agent: dynamic data, such as rapidly changing weather metrics and unpredictable user profiles, must be instantly accessible to the Node.js execution engine without the latency overhead of complex SQL JOIN operations.")
    doc.add_paragraph("The primary logical stores include:")
    doc.add_paragraph("• User Profile Store: Holds the user's physiological preferences (activity type, health sensitivity, location) used by the Prediction Engine.")
    doc.add_paragraph("• Environmental Context Store: Holds the nested, hierarchical weather data (Temperature, Humidity, AQI metrics) fetched from external APIs.")
    doc.add_paragraph("• Prediction Cache Store: A temporal storage system that caches the AI-generated health heuristics to prevent redundant recalculations and establish historical safety baselines.")

    doc.add_heading('3.4.2 Physical Storage Design', level=3)
    doc.add_paragraph("The physical database is implemented on a distributed MongoDB Atlas cluster, orchestrated via the Mongoose Object Data Modeling (ODM) library within the Node.js backend.")
    doc.add_paragraph("Instead of SQL CREATE TABLE commands, the schema is enforced programmatically via Mongoose Schemas. For example, the primary PredictionLog data store is structured with the following JSON/BSON schema logic:")
    
    json_code = '''{
    "userProfile": {
        "activity_type": "Running",
        "health_sensitivity": "Asthma",
        "city": "Indore"
    },
    "environmentalData": {
        "temperature_c": 34.5,
        "humidity_percent": 65,
        "pm25_concentration": 112.4
    },
    "heuristics": {
        "heat_index_c": 41.2,
        "hydration_ml_hr": 850,
        "lung_recovery_hours": 14,
        "safety_level": "DANGER"
    },
    "timestamp": "2026-04-28T10:30:00Z"
}'''
    doc.add_paragraph(json_code, style='Intense Quote')
    
    doc.add_paragraph("By utilizing this physical NoSQL architecture, the system guarantees thread-safe, low-latency I/O operations from the Express server. Furthermore, sensitive credentials required by the External API Fetchers and the MongoDB cluster URI are physically isolated in a local .env file, ensuring that API keys and connection strings are never accidentally serialized into the database or pushed to version control.")

    output_path = r'Report\AtmoSense_Ch1_to_Ch3_Persistence_Final.docx'
    doc.save(output_path)
    print("Persistence design section successfully appended to the document.")

if __name__ == '__main__':
    add_persistence()
