import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_front_pages():
    doc = docx.Document()
    
    # --- Report Approval ---
    doc.add_heading('Report Approval', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("The project work \"AtmoSense AI - Smart Environmental Health Intelligence\" is hereby approved as a creditable study of an engineering/computer application subject carried out and presented in a manner satisfactory to warrant its acceptance as a prerequisite for the Degree for which it has been submitted.")
    doc.add_paragraph("It is to be understood that by this approval the undersigned do not endorse or approve any statement made, opinion expressed, or conclusion drawn therein; but approve the \"Project Report\" only for the purpose for which it has been submitted.\n\n")
    
    doc.add_paragraph("Internal Examiner\nName:\nDesignation:\nAffiliation:\n\n")
    doc.add_paragraph("External Examiner\nName:\nDesignation:\nAffiliation:")
    doc.add_page_break()
    
    # --- Declaration ---
    doc.add_heading('Declaration', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("We hereby declare that the project entitled \"AtmoSense AI - Smart Environmental Health Intelligence\" submitted in partial fulfillment for the award of the degree of Bachelor of Technology in Computer Science & Engineering, completed under the supervision of Prof. Vishal Sharma and Prof. Suyog Munshi, Faculty of Engineering, Medicaps University Indore, is an authentic work.")
    doc.add_paragraph("Further, we declare that the content of this Project work, in full or in parts, have neither been taken from any other source nor have been submitted to any other Institute or University for the award of any degree or diploma.\n\n")
    doc.add_paragraph("Signature and name of the student(s) with date:\n\n")
    doc.add_paragraph("________________________________\nPiyush Solanki (EN23CS304044)\n\n________________________________\nMayank Sahu (EN23CS304037)\n\n________________________________\nPiyush Pal (EN23CS304044)\n\nDate: _______________")
    doc.add_page_break()
    
    # --- Certificate ---
    doc.add_heading('Certificate', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("We, Prof. Vishal Sharma and Prof. Suyog Munshi, certify that the project entitled \"AtmoSense AI - Smart Environmental Health Intelligence\" submitted in partial fulfillment for the award of the degree of Bachelor of Technology in Computer Science & Engineering by Piyush Solanki, Mayank Sahu, Piyush Pal is a record of the work carried out by them under our guidance and that the work has not formed the basis of award of any other degree elsewhere.\n\n\n")
    doc.add_paragraph("________________________________\nProf. Vishal Sharma \nDepartment of Computer Science & Engineering\nMedicaps University, Indore\n\n")
    doc.add_paragraph("________________________________\nProf. Suyog Munshi \nDepartment of Computer Science & Engineering\nMedicaps University, Indore\n\n")
    doc.add_paragraph("________________________________\nDr. Kailash Bandhu\nHead of the Department\nComputer Science & Engineering\nMedicaps University, Indore")
    doc.add_page_break()
    
    # --- Acknowledgements ---
    doc.add_heading('Acknowledgements', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("We would like to express our deepest gratitude to the Honorable Chancellor, Shri R C Mittal, who has provided us with every facility to successfully carry out this project, and our profound indebtedness to Prof. (Dr.) D. K. Patnaik, Vice Chancellor, Medicaps University, whose unfailing support and enthusiasm has always boosted our morale. We also thank Prof. (Dr.) Ratnesh Litoriya, Dean, Faculty of Engineering, Medicaps University, for giving us a chance to work on this project. We would also like to thank our Head of the Department Prof. (Dr.) Kailash Bandhu for his continuous encouragement for the betterment of the project.")
    doc.add_paragraph("We are immensely grateful to our project guides, Prof. Vishal Sharma and Prof. Suyog Munshi, for their invaluable guidance, constant motivation, and expert advice throughout the course of this project. Their insights greatly shaped the direction and quality of our work.")
    doc.add_paragraph("We extend our sincere thanks to all faculty members and staff of the Department of Computer Science & Engineering for their continuous support and cooperation. Without their help and support, this report would not have been possible.\n\n")
    doc.add_paragraph("Piyush Solanki\nMayank Sahu\nPiyush Pal\nB.Tech. III Year\nDepartment of Computer Science & Engineering\nFaculty of Engineering\nMedicaps University, Indore")
    doc.add_page_break()
    
    # --- Abstract ---
    doc.add_heading('Abstract', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("AtmoSense AI is a modular, AI-powered environmental intelligence ecosystem designed to assist individuals and professionals in autonomously monitoring complex meteorological and physiological workflows. In today's highly fragmented application environment, executing tasks such as tracking real-time AQI, predicting hydration needs, and assessing heat index requires immense cognitive load and constant context-switching between disparate health and weather apps. This project proposes and implements a comprehensive web-based architecture that bridges human intent with autonomous health heuristic execution. The system leverages a highly resilient API orchestration layer—integrating OpenWeatherMap and WAQI APIs—to ensure continuous availability and mitigate API rate-limits via a proprietary Mock Data Fallback mechanism. Built on a robust MERN stack (MongoDB, Express.js, React.js, Node.js), the application interfaces securely via REST APIs. Users can view real-time data, submit health profiles, calculate lung recovery times, and automatically persist prediction logs to MongoDB Atlas entirely from a unified Dashboard. This eliminates the need to manually navigate disparate weather websites and manual health calculators. The system significantly reduces operational friction, ensuring zero downtime during network disruptions while aggregating personalized physiological intelligence based on the user's physical activity. Testing results demonstrate that the engine executes multi-step fetching and heuristic computation in milliseconds, with exceptionally high accuracy. AtmoSense AI addresses the critical need for a deeply integrated environmental assistant by providing a secure, scalable, and responsive solution that radically enhances daily health planning and outdoor safety.")
    doc.add_paragraph("Keywords: Environmental Intelligence, Health Heuristics, React.js, Node.js, MERN Stack, Mock Data Fallback, API Orchestration, Data Persistence, Weather Tracking.")
    doc.add_page_break()
    
    # --- Table of Contents (Simulated) ---
    doc.add_heading('Table of Contents', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = """Report Approval	ii
Declaration	iii
Certificate	iv
Acknowledgements	v
Abstract	vi
Table of Contents	vii
List of Figures	viii
List of Tables	ix
Abbreviations	x
Chapter 1	  Introduction	1
	1.1  Introduction	1
	1.2  Literature Review	2
	1.3  Objectives	3
	1.4  Significance	3
	1.5  Research Design / Methodology	4
	1.6  Source of Data	4
	1.7  Chapter Scheme	5
Chapter 2	  Requirements Specification	6
	2.1  User Characteristics	6
	2.2  Functional Requirements	6
	2.3  Dependencies	7
	2.4  Performance Requirements	8
	2.5  Hardware Requirements	8
	2.6  Constraints and Assumptions	8
Chapter 3	  Design	9
	3.1  Algorithm	9
	3.2  Function Oriented Design	10
	3.3  System Design (DFD, Flow, Sequence, etc.)	10
	3.4  NoSQL Data Persistence Design	12
Chapter 4	  Implementation, Testing, and Maintenance	13
	4.1  Languages, IDEs, Tools and Technologies	13
	4.2  Testing Techniques and Test Plans	14
	4.3  Installation Instructions	15
	4.4  End User Instructions	15
Chapter 5	  Results and Discussions	16
	5.1  User Interface Representation	16
	5.2  Brief Description of Modules	16
    5.3  System Operations: Snapshots  17
Chapter 6	  Summary and Conclusions	18
Chapter 7	  Future Scope	19"""
    doc.add_paragraph(toc)
    doc.add_page_break()
    
    # --- List of Figures ---
    doc.add_heading('List of Figures', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    figures = """Figure 3.1   Data Flow Diagram – Level 0 (Context Diagram)
Figure 3.2   Data Flow Diagram – Level 1
Figure 3.3   Activity Diagram of AtmoSense AI
Figure 3.4   System Flow Chart (Mock Data Fallback)
Figure 3.5   Class Diagram (MERN Architecture)
Figure 3.6   Entity-Relationship (ER) Diagram (MongoDB Schema)
Figure 3.7   Sequence Diagram
Figure 5.1   Dashboard Screenshot
Figure 5.2   Health Prediction Output Screenshot"""
    doc.add_paragraph(figures)
    doc.add_page_break()
    
    # --- List of Tables ---
    doc.add_heading('List of Tables', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    tables = """Table 2.1   Functional Requirements of AtmoSense AI
Table 2.2   Hardware & Software Requirements
Table 4.1   Test Cases for Core Operations (API Fallback, Heuristics, Persistence)"""
    doc.add_paragraph(tables)
    doc.add_page_break()
    
    # --- Abbreviations ---
    doc.add_heading('Abbreviations', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    abbrev = """AI         –  Artificial Intelligence
API        –  Application Programming Interface
AQI        –  Air Quality Index
BSON       –  Binary JSON
CRUD       –  Create, Read, Update, Delete
DB         –  Database
DFD        –  Data Flow Diagram
ER         –  Entity-Relationship
HTTP       –  Hypertext Transfer Protocol
IDE        –  Integrated Development Environment
JSON       –  JavaScript Object Notation
MERN       –  MongoDB, Express.js, React.js, Node.js
REST       –  Representational State Transfer
SPA        –  Single Page Application
UI         –  User Interface"""
    doc.add_paragraph(abbrev)

    output_path = r'Report\AtmoSense_Front_Pages.docx'
    doc.save(output_path)
    print("Front pages document generated.")

if __name__ == '__main__':
    create_front_pages()
