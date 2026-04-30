import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def create_detailed_report():
    doc = docx.Document()
    
    # --- Title Page ---
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('AtmoSense AI\n(Smart Environmental Health Intelligence)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(24)
        
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('A Minor Project Report\nSubmitted in partial fulfillment of requirement of the\nDegree of\nBACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING\nBY')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Piyush Solanki (EN23CS304044)\nMayank Sahu (EN23CS304037)\nPiyush Pal (EN23CS304044)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Under the Guidance of\nProf. Vishal Sharma\nProf. Suyog Munshi')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Department of Computer Science & Engineering\nFaculty of Engineering\nMEDICAPS UNIVERSITY, INDORE- 453331\n\nAPRIL-2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract_text = """The AtmoSense AI (Smart Environmental Health Intelligence) is a comprehensive, full-stack web ecosystem engineered to collect real-time meteorological and air quality data, subsequently transforming it into highly personalized health and outdoor activity insights. As environmental factors like extreme temperatures and elevated pollution levels increasingly pose significant health risks, there is a pressing demand for intelligent monitoring systems capable of proactive health risk assessment. AtmoSense AI addresses this challenge by providing an interactive, user-centric dashboard that aggregates live data including temperature, humidity, wind speed, UV index, and various Air Quality Index (AQI) pollutants (PM2.5, PM10, NO2, O3). 

Beyond mere data visualization, the core innovation of AtmoSense AI lies within its proprietary AI Health Prediction Engine. This engine employs complex algorithms and heuristics, such as the Rothfusz regression for Heat Index, specialized hydration requirement estimators, and lung recovery time calculators conditioned on individual health sensitivities (e.g., asthma). By synthesizing this data, the platform issues color-coded safety recommendations to guide users' outdoor activities. 

Built upon the robust MERN (MongoDB, Express.js, React.js, Node.js) stack, the application features a modern, responsive glassmorphism UI developed with Tailwind CSS and Recharts. The backend seamlessly integrates with external APIs like OpenWeatherMap and WAQI, featuring an intelligent mock-data fallback mechanism to ensure uninterrupted service. This report details the comprehensive lifecycle of AtmoSense AI, encompassing requirement analysis, system design, architectural workflows, implementation details, and future potential."""
    doc.add_paragraph(abstract_text)
    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc = """1. Introduction
2. Literature Review
3. System Requirements & Specification
4. System Architecture & Design
5. Workflow Diagrams & Data Flow
6. Implementation Details & Technologies
7. AI Prediction Engine Algorithms
8. Testing & Validation
9. Results & Snapshots
10. Conclusion & Future Scope
"""
    doc.add_paragraph(toc)
    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph("In recent years, the detrimental effects of climate change and rapid urbanization have led to severe fluctuations in weather patterns and alarming increases in air pollution. Poor air quality, specifically high concentrations of particulate matter (PM2.5 and PM10) and ground-level ozone, is directly linked to respiratory and cardiovascular diseases. Similarly, extreme heat events pose significant risks of heatstroke and severe dehydration, especially during outdoor physical activities. Traditional weather applications provide raw meteorological data but fail to translate these numbers into actionable, personalized health advice. There is a critical gap between environmental data availability and individual health risk management.")
    doc.add_paragraph("AtmoSense AI bridges this gap by acting as a smart intermediary. It not only fetches real-time environmental data but also processes it through an AI-driven health prediction engine to output personalized recommendations. Whether a user is planning a rigorous outdoor run or a casual commute, AtmoSense AI calculates the precise impact of the current atmosphere on their body.")
    
    doc.add_heading('1.2 Objectives', level=2)
    doc.add_paragraph("The primary objectives of the AtmoSense AI project are:")
    doc.add_paragraph("1. Real-Time Data Aggregation: To develop a robust backend system capable of concurrently fetching data from multiple external sources (OpenWeatherMap for meteorology and WAQI for air quality).")
    doc.add_paragraph("2. Personalized Health Inference: To design and implement algorithms that calculate Heat Index, estimate core body temperature elevation, compute activity-specific hydration needs, and predict lung recovery times based on user-specific health conditions (e.g., Asthma).")
    doc.add_paragraph("3. Intuitive User Interface: To construct a highly responsive, aesthetically pleasing frontend utilizing glassmorphism principles, ensuring data is accessible and visually engaging.")
    doc.add_paragraph("4. System Resilience: To implement a fallback mechanism that generates realistic mock data when external APIs are unreachable, ensuring continuous application availability.")
    
    doc.add_heading('1.3 Scope of the Project', level=2)
    doc.add_paragraph("The scope encompasses the development of a web-based application accessible via modern desktop and mobile browsers. It includes user input forms for profiling (city, activity type, health sensitivity), interactive dashboards for data visualization, and a persistent database to log prediction histories. The project currently focuses on immediate, real-time predictions rather than long-term historical forecasting.")
    doc.add_page_break()

    # --- Chapter 2: Literature Review ---
    doc.add_heading('Chapter 2: Literature Review', level=1)
    doc.add_paragraph("The intersection of environmental science, health informatics, and software engineering has seen significant research. Studies by the World Health Organization (WHO) emphasize the critical need for real-time air quality monitoring to mitigate health risks. Existing platforms like AQICN provide extensive pollution data, while applications like AccuWeather offer localized weather forecasts.")
    doc.add_paragraph("However, literature indicates a lack of integrated platforms that synthesize both datasets to produce personalized health heuristics. Most tools require the user to manually interpret what an AQI of 150 or a temperature of 35°C means for their specific planned activity. Research into heat stress indices (such as the Rothfusz regression model used by the National Weather Service) and hydration physiology highlights the complex variables involved in human thermoregulation.")
    doc.add_paragraph("AtmoSense AI builds upon these established scientific models, codifying them into automated algorithms within a scalable web architecture, thereby advancing the practical application of environmental health informatics.")
    doc.add_page_break()

    # --- Chapter 3: System Requirements & Specification ---
    doc.add_heading('Chapter 3: System Requirements & Specification', level=1)
    
    doc.add_heading('3.1 Functional Requirements', level=2)
    doc.add_paragraph("- The system must allow users to search for real-time weather and AQI by city name.")
    doc.add_paragraph("- The system must present data visually using charts and numerical cards.")
    doc.add_paragraph("- The system must accept user profiles consisting of activity type (e.g., Running, Gym, Walking) and health sensitivity (e.g., Asthma, Normal).")
    doc.add_paragraph("- The system must generate a comprehensive health report including Heat Index, Temperature Increase, Hydration Requirement, Lung Recovery Time, and a Safety Recommendation.")
    doc.add_paragraph("- The system must store past predictions in a MongoDB database for historical logging.")
    
    doc.add_heading('3.2 Non-Functional Requirements', level=2)
    doc.add_paragraph("- Performance: API responses should be processed and rendered within 2 seconds.")
    doc.add_paragraph("- Reliability: The system must fall back to mock data gracefully if external APIs fail.")
    doc.add_paragraph("- Usability: The UI must be intuitive, featuring clear color-coded indicators (Red, Yellow, Green) for safety levels.")
    doc.add_paragraph("- Scalability: The backend architecture must support asynchronous handling of multiple concurrent user requests.")
    
    doc.add_heading('3.3 Hardware & Software Requirements', level=2)
    doc.add_paragraph("Software:\n- Operating System: Windows / Linux / macOS\n- Environment: Node.js (v18+)\n- Database: MongoDB Atlas or Local instance\n- Frontend Framework: React.js via Vite\n- Package Manager: npm or yarn\n\nHardware (Minimum for Server):\n- CPU: Dual-core 2.0 GHz\n- RAM: 4 GB\n- Storage: 10 GB available space")
    doc.add_page_break()

    # --- Chapter 4: System Architecture & Design ---
    doc.add_heading('Chapter 4: System Architecture & Design', level=1)
    doc.add_paragraph("AtmoSense AI utilizes a decoupled Client-Server architecture based on the MERN stack.")
    
    doc.add_heading('4.1 High-Level Architecture', level=2)
    doc.add_paragraph("The architecture comprises three main tiers:")
    doc.add_paragraph("1. Presentation Tier (Frontend): Built with React and Tailwind CSS. It manages the user interface, state management, routing, and asynchronous API calls to the backend using Axios. Key components include the main Dashboard, Prediction Form, and Result Visualizations.")
    doc.add_paragraph("2. Logic Tier (Backend): An Express.js server running on Node.js. It acts as a middleware, orchestrating requests to external APIs, processing the raw data through the AI Prediction Engine, and formatting the response. It handles business logic, data normalization, and error handling.")
    doc.add_paragraph("3. Data Tier (Database): MongoDB handles the persistence layer, storing 'PredictionLog' documents that contain the user's input parameters alongside the fetched environmental data and the calculated health impacts.")
    
    doc.add_heading('4.2 Architectural Diagram Description', level=2)
    doc.add_paragraph("The User interacts with the React Frontend (Browser). The Frontend dispatches HTTP GET/POST requests to the Express.js Server. The Server's Services Layer concurrently fetches data from OpenWeatherMap and WAQI APIs. Upon receiving the data, it is passed to the Prediction Engine module for heuristic analysis. Finally, the server logs the transaction in MongoDB and returns the comprehensive JSON payload to the Frontend for rendering.")
    doc.add_page_break()

    # --- Chapter 5: Workflow Diagrams & Data Flow ---
    doc.add_heading('Chapter 5: Workflow Diagrams & Data Flow', level=1)
    
    doc.add_heading('5.1 Complete User Journey', level=2)
    doc.add_paragraph("The user journey is meticulously designed for seamless interaction. Upon opening the application, the user lands on the Home Dashboard. Here, a default city's weather and AQI are displayed. The user can navigate to the 'Predict' page to access the AI features.")
    doc.add_paragraph("In the Prediction Flow:\n1. The user enters a target city location.\n2. The user selects an Activity Type (Walking, Running, Gym, Commuting).\n3. The user selects a Health Sensitivity profile (Normal, Asthma, Heat Sensitive).\n4. Upon submission, the frontend displays a loading state while calling the backend.\n5. The backend returns the analyzed data, and the user is navigated to the Results Dashboard featuring safety badges, heat index metrics, and lung recovery progress bars.")
    
    doc.add_heading('5.2 Backend API Pipeline (Sequence)', level=2)
    doc.add_paragraph("When a POST request is made to '/predict':\n- The Express server extracts the payload.\n- It triggers `fetchWeather()` and `fetchAqi()` in parallel.\n- If API keys are valid, external data is fetched; otherwise, the Mock Fallback generates realistic synthetic data.\n- The aggregated data is piped into `runPredictions(weather, aqi, userInput)`.\n- The output is structured, saved to MongoDB via Mongoose, and returned to the client with a 200 OK status.")
    
    doc.add_heading('5.3 AI Prediction Engine Workflow', level=2)
    doc.add_paragraph("The internal logic of the prediction engine involves several discrete sub-modules:")
    doc.add_paragraph("- Heat Index Calculation: Evaluates if temperature < 80°F (uses simple formula) or > 80°F (uses Rothfusz Regression).")
    doc.add_paragraph("- Body Temp Increase: Adds a base temperature based on activity (e.g., Running = +0.8°C) and scales it if the Heat Index exceeds 32°C.")
    doc.add_paragraph("- Hydration Calculation: Starts at a baseline of 200 ml/hr, adds volume based on activity intensity, and applies a heat multiplier if ambient temp > 25°C.")
    doc.add_paragraph("- Lung Recovery: Starts at 1.0 hr, adds time based on AQI severity (> 50), and applies a 1.5x multiplier if the user is Asthma sensitive.")
    doc.add_paragraph("- Safety Recommendation: Evaluates all thresholds to classify the condition as RED (Dangerous), YELLOW (Moderate), or GREEN (Safe).")
    doc.add_page_break()

    # --- Chapter 6: Implementation Details ---
    doc.add_heading('Chapter 6: Implementation Details & Technologies', level=1)
    
    doc.add_heading('6.1 Frontend Technologies', level=2)
    doc.add_paragraph("React.js & Vite: Chosen for its rapid development server and component-based architecture. Vite provides significantly faster Hot Module Replacement (HMR) compared to Create React App.")
    doc.add_paragraph("Tailwind CSS: Utilized for utility-first styling, enabling the rapid creation of the glassmorphism aesthetic without writing extensive custom CSS files. It facilitated responsive design through arbitrary breakpoints.")
    doc.add_paragraph("Framer Motion: Implemented to add fluid, physics-based animations to page transitions, modal appearances, and background elements, greatly enhancing the premium feel of the application.")
    doc.add_paragraph("Recharts: A composable charting library built on React components, used to render the multi-pollutant bar charts on the Air Quality dashboard.")
    
    doc.add_heading('6.2 Backend Technologies', level=2)
    doc.add_paragraph("Node.js & Express: Provide a lightweight, non-blocking, event-driven runtime suitable for handling numerous concurrent API requests.")
    doc.add_paragraph("Axios: Used within the backend services layer to execute HTTP requests to external APIs due to its built-in promise support and automatic JSON data transformation.")
    doc.add_paragraph("MongoDB & Mongoose: MongoDB offers a flexible schema-less document store perfect for saving heterogeneous API responses. Mongoose provides a rigorous modeling environment, ensuring data integrity before persistence.")
    
    doc.add_heading('6.3 Core Code Snippets Insight', level=2)
    doc.add_paragraph("The application heavily relies on asynchronous JavaScript. In the backend `index.js`, the `/predict` route utilizes `Promise.allSettled` or sequential `await` calls to gather data securely. The `calculateImpact` module acts as a pure function, taking state as arguments and returning deterministic calculations, ensuring testability and reliability.")
    doc.add_page_break()

    # --- Chapter 7: AI Algorithms ---
    doc.add_heading('Chapter 7: AI Prediction Engine Algorithms Detailed', level=1)
    
    doc.add_heading('7.1 The Rothfusz Regression (Heat Index)', level=2)
    doc.add_paragraph("The Heat Index represents how hot it really feels when relative humidity is factored in with the actual air temperature. AtmoSense AI implements the Rothfusz regression equation:")
    doc.add_paragraph("HI = -42.379 + 2.04901523*T + 10.14333127*R - 0.22475541*T*R - 0.00683783*T^2 - 0.05481717*R^2 + 0.00122874*T^2*R + 0.00085282*T*R^2 - 0.00000199*T^2*R^2")
    doc.add_paragraph("Where T is temperature in Fahrenheit and R is relative humidity percentage. The algorithm automatically handles unit conversions from Celsius and applies the necessary adjustments if the conditions meet specific high-heat/low-humidity criteria.")
    
    doc.add_heading('7.2 Lung Recovery Algorithm', level=2)
    doc.add_paragraph("This heuristic estimates physiological recovery time post-exposure to pollutants. It is formulated as:")
    doc.add_paragraph("Base Time = 1.0 hours")
    doc.add_paragraph("If AQI > 50: Additional Time = (AQI - 50) * 0.02 hours")
    doc.add_paragraph("If User Profile == 'Asthma Sensitive': Total Time = Total Time * 1.5")
    doc.add_paragraph("This linear interpolation approach provides a practical, conservative estimate to warn sensitive users of prolonged respiratory stress.")
    doc.add_page_break()

    # --- Chapter 8: Testing ---
    doc.add_heading('Chapter 8: Testing & Validation', level=1)
    doc.add_paragraph("To ensure system robustness, rigorous testing methodologies were employed.")
    
    doc.add_heading('8.1 Unit Testing', level=2)
    doc.add_paragraph("The pure functions within the AI Prediction Engine were isolated and tested with boundary values. For instance, the Heat Index calculator was tested with extremely high humidity (99%) and low temperature (10°C) to ensure the fallback 'simple formula' executed correctly without returning anomalous high values.")
    
    doc.add_heading('8.2 Integration Testing', level=2)
    doc.add_paragraph("The interaction between the Express routes and the external APIs was validated. Test cases simulated scenarios where the external API timed out or returned 404 errors, ensuring the backend successfully caught the exceptions and triggered the internal Mock Fallback generator instead of crashing.")
    
    doc.add_heading('8.3 System Testing', level=2)
    doc.add_paragraph("End-to-end testing was conducted manually via the React UI. Users navigating from the Home page, inputting data into the Predict form, and viewing the Results page were traced to ensure state was correctly passed via React Router's location state and that no memory leaks occurred during component unmounting.")
    doc.add_page_break()

    # --- Chapter 9: Results & Snapshots ---
    doc.add_heading('Chapter 9: Results & Snapshots', level=1)
    doc.add_paragraph("The AtmoSense AI platform successfully delivers a cohesive user experience. The integration of modern design paradigms (glassmorphism, dark mode aesthetics, dynamic particle backgrounds) combined with rapid backend responses resulted in a highly performant application.")
    
    doc.add_heading('9.1 Dashboard Interface', level=2)
    doc.add_paragraph("The primary dashboard displays the live AQI via a custom SVG circular gauge that dynamically colors itself (Green to Purple) based on the severity of the air quality. Parallel cards display temperature, humidity, and wind speed with crisp Lucide icons.")
    
    doc.add_heading('9.2 Prediction Results Interface', level=2)
    doc.add_paragraph("Upon running a prediction, users are presented with actionable cards. A prominent colored banner dictates the overall safety recommendation (e.g., 'YELLOW: Moderate conditions, reduce intensity'). Sub-cards detail the exact ml/hr of water to drink and the hours required for lung recovery, accompanied by visual progress bars.")
    doc.add_paragraph("[Note: In the final printed report, actual screenshots of the UI should be inserted here to demonstrate the visual fidelity of the application.]")
    doc.add_page_break()

    # --- Chapter 10: Conclusion ---
    doc.add_heading('Chapter 10: Conclusion & Future Scope', level=1)
    
    doc.add_heading('10.1 Conclusion', level=2)
    doc.add_paragraph("The AtmoSense AI project has successfully demonstrated the feasibility and immense value of synthesizing real-time environmental data with AI-driven health heuristics. By abstracting complex meteorological and pollutant metrics into easily digestible, personalized safety recommendations, the platform empowers users to make informed decisions regarding their outdoor activities. The robust MERN architecture, coupled with resilient fallback mechanisms and a premium user interface, ensures that the application is both reliable and engaging.")
    
    doc.add_heading('10.2 Future Enhancements', level=2)
    doc.add_paragraph("While the current iteration of AtmoSense AI is highly functional, several avenues for future development exist:")
    doc.add_paragraph("1. Machine Learning Integration: Transitioning from heuristic-based algorithms to true machine learning models (e.g., Random Forests or Neural Networks) trained on vast datasets of historical health and weather data to provide even more accurate predictions.")
    doc.add_paragraph("2. Wearable Device Sync: Integrating with APIs from smartwatches (Apple Watch, Garmin, Fitbit) to pull real-time biometric data (heart rate, blood oxygen) to calibrate the prediction engine dynamically during the activity.")
    doc.add_paragraph("3. Mobile Application: Developing a native mobile application using React Native to utilize push notifications, alerting users proactively when air quality drops in their current GPS location.")
    doc.add_paragraph("4. Multi-Day Forecasting: Extending the prediction engine to analyze 5-day weather forecasts to help users plan their workout schedules for the entire week.")
    doc.add_paragraph("\n\n\n\n\n[End of Report]")
    
    # Save the huge document
    doc.save(r'Report\AtmoSense_Detailed_Project_Report.docx')

if __name__ == '__main__':
    create_detailed_report()
