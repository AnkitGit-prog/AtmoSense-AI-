import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title Page
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('AtmoSense AI\n(Smart Environmental Health Intelligence)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = docx.shared.Pt(24)
        
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('A Minor Project Report\nSubmitted in partial fulfillment of requirement of the\nDegree of\nBACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING\nBY')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Piyush Solanki\nEN23CS304044\n\nMayank Sahu\nEN23CS304037\n\nPiyush Pal\nEN23CS304044')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Under the Guidance of\nProf. Vishal Sharma\nProf. Suyog Munshi')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Department of Computer Science & Engineering\nFaculty of Engineering\nMEDICAPS UNIVERSITY, INDORE- 453331\n\nAPRIL-2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph("AtmoSense AI is a full-stack, AI-driven web application that collects real-time weather and Air Quality data, converting it into personalized health and outdoor activity insights. The system provides a unified dashboard for viewing live temperature, humidity, wind speed, UV index, and AQI (Air Quality Index). It leverages an AI health prediction engine to calculate heat index, hydration requirements, lung recovery time, and safety recommendations based on individual user profiles. Built with a modern technology stack including React.js, Node.js, Express, and MongoDB, AtmoSense AI aims to help users make informed decisions about their outdoor activities and health in changing environmental conditions.")
    
    doc.add_paragraph("Keywords: AtmoSense AI, Environmental Health, AQI, Weather Tracking, AI Health Prediction, React, Node.js, MERN Stack.")

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('Chapter 1: Introduction', level=1)
    
    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph("Environmental factors such as air quality and extreme weather conditions significantly impact human health. With the rise of pollution and climate change, there is a growing need for smart systems that can analyze these factors in real-time and provide actionable health insights. AtmoSense AI is designed to address this need by offering a comprehensive platform that monitors weather and air quality metrics and uses an AI engine to predict personal health impacts, such as required hydration and lung recovery times.")
    
    doc.add_heading('1.2 Objectives', level=2)
    doc.add_paragraph("1. To develop a real-time dashboard for monitoring weather and air quality.\n2. To implement an AI-driven prediction engine that calculates personal health impacts based on environmental data.\n3. To provide actionable safety recommendations for outdoor activities.\n4. To build a robust, scalable full-stack application using the MERN stack.")
    
    doc.add_heading('1.3 Significance', level=2)
    doc.add_paragraph("The significance of AtmoSense AI lies in its proactive approach to environmental health. By transforming raw meteorological and pollution data into personalized, easy-to-understand health metrics, it empowers users to protect their well-being. The system's ability to factor in individual health sensitivities, such as asthma, makes it a valuable tool for vulnerable populations.")
    
    doc.add_heading('1.4 Methodology', level=2)
    doc.add_paragraph("The development of AtmoSense AI followed a structured methodology:\n- Requirement Analysis: Defining the necessary environmental parameters and health metrics.\n- System Design: Architecting the frontend with React and Tailwind CSS, and the backend with Node.js and MongoDB.\n- Implementation: Developing the API integrations for OpenWeatherMap and WAQI, and constructing the AI prediction algorithms.\n- Testing: Validating the prediction logic and ensuring seamless data flow between the frontend and backend.\n- Deployment: Setting up the environment for local and potential cloud deployment.")

    doc.add_page_break()

    # Chapter 2: System Design
    doc.add_heading('Chapter 2: System Design', level=1)
    doc.add_paragraph("The system is built on a client-server architecture using the MERN stack.")
    
    doc.add_heading('2.1 Frontend Architecture', level=2)
    doc.add_paragraph("The frontend is developed using React.js (via Vite) for high performance. It features a modern glassmorphism UI styled with Tailwind CSS. Data visualization is handled by Recharts, providing interactive charts for pollutant levels. Key components include the Home dashboard, Input Form for predictions, and the Air Quality page.")
    
    doc.add_heading('2.2 Backend Architecture', level=2)
    doc.add_paragraph("The backend is powered by Node.js and Express.js, providing RESTful API endpoints. It integrates with external APIs (OpenWeatherMap and WAQI) to fetch real-time data, with a built-in mock data fallback mechanism. The MongoDB database, interfaced via Mongoose, stores prediction logs.")
    
    doc.add_heading('2.3 AI Prediction Engine', level=2)
    doc.add_paragraph("The core logic involves several algorithms:\n- Heat Index Calculation: Uses temperature and humidity to estimate the felt temperature.\n- Hydration Alert: Calculates recommended water intake based on activity intensity and environmental conditions.\n- Lung Recovery Time: Estimates recovery time based on PM2.5 levels and health sensitivity.\n- Safety Recommendation: Provides color-coded alerts (Green, Yellow, Red).")

    doc.add_page_break()

    # Chapter 3: Implementation
    doc.add_heading('Chapter 3: Implementation', level=1)
    doc.add_paragraph("The implementation phase involved coding the defined architecture into a functional system.")
    
    doc.add_heading('3.1 Technologies Used', level=2)
    doc.add_paragraph("- Frontend: React.js, Vite, Tailwind CSS, Framer Motion, Recharts, Lucide React.\n- Backend: Node.js, Express.js, Axios, Mongoose.\n- Database: MongoDB Atlas (with local memory server fallback).")
    
    doc.add_heading('3.2 Key Features', level=2)
    doc.add_paragraph("- Live Weather & AQI Dashboard.\n- AI Health Prediction Engine.\n- Responsive, animated UI.")

    doc.add_page_break()

    # Chapter 4: Conclusion and Future Scope
    doc.add_heading('Chapter 4: Conclusion & Future Scope', level=1)
    
    doc.add_heading('4.1 Conclusion', level=2)
    doc.add_paragraph("AtmoSense AI successfully bridges the gap between raw environmental data and personal health awareness. By combining real-time API data with customized prediction algorithms, the system provides users with valuable insights into how weather and air quality affect their bodies during outdoor activities.")
    
    doc.add_heading('4.2 Future Scope', level=2)
    doc.add_paragraph("Future enhancements may include:\n- Mobile application development for on-the-go access.\n- Integration with wearable devices (e.g., smartwatches) for real-time biometric feedback.\n- Advanced machine learning models for more precise health impact forecasting based on historical user data.")

    doc.save(r'Report\AtmoSense_Mini_Project_Report.docx')
    print("Report generated successfully.")

if __name__ == '__main__':
    create_report()

