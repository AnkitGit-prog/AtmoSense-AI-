import base64
import urllib.request
import docx
from docx.shared import Inches
import os

def add_activity_image():
    # Mermaid diagram for AtmoSense Activity Flow
    mermaid_code = """
flowchart TD
    User([User])
    User --> Dash[View Dashboard Live Weather and AQI]
    User --> Pred[Predict Health Impact Select Activity]
    User --> AQI[Analyze Air Quality Pollutant Breakdown]
    
    Pred --> B[Backend API Parallel Fetch]
    B --> C{API Keys Valid?}
    C -- Yes --> D[Run AI Prediction Algorithms]
    C -- No --> E[Fallback to Mock Data Generator]
    E --> D
    
    D --> F[Persist to MongoDB]
    F --> G[Render Results & Safety Badges on React UI]
    """
    
    b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}?type=png&bgColor=!white"
    
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0'}
    )
    
    img_path = r'Report\Activity_Diagram.png'
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as out_file:
                out_file.write(response.read())
        print("Activity Diagram image downloaded successfully.")
    except Exception as e:
        print(f"Failed to download image: {e}")
        return

    # Insert image into document
    doc_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Complete.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_paragraph('\nActivity Diagram Image:')
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph('\n')

    output_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Final.docx'
    doc.save(output_path)
    print("Activity diagram image appended to the document.")

if __name__ == '__main__':
    add_activity_image()
