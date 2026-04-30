import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

def append_chapter_3():
    # Load the existing document
    doc_path = r'Report\AtmoSense_Ch1_Ch2_Detailed_v2.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find v2 document.")
        return
        
    doc = docx.Document(doc_path)
    
    doc.add_page_break()

    # --- Chapter 3 ---
    doc.add_heading('Chapter 3\nDesign', level=1)
    
    doc.add_heading('3.1 Algorithm', level=2)
    doc.add_paragraph("The core algorithm implemented in AtmoSense AI follows a strictly defined, multi-stage pipeline for intent parsing, environmental context gathering, and heuristic prediction execution. This process is governed by the proprietary 'Mock Data Fallback' and 'Concurrent API Aggregation' mechanisms. The algorithm is described below:")
    doc.add_paragraph("Step 1 – User Intent & Context Extraction: The React frontend receives user input. The system extracts the necessary location context and physiological parameters (City, Activity Type, Health Sensitivity) via controlled form components and state management.")
    doc.add_paragraph("Step 2 – HTTP Request Construction & Orchestration: The extracted data is sanitized to prevent injection attacks and dispatched asynchronously as a structured POST request to the Node.js / Express.js backend API.")
    doc.add_paragraph("Step 3 – Primary API Routing (Parallel Fetch): The backend Node.js execution engine initiates asynchronous REST API calls concurrently using Promise.allSettled to the primary configured environmental providers (OpenWeatherMap for meteorology and WAQI for air quality).")
    doc.add_paragraph("Step 4 – Mock Data Fallback Interception: The system actively listens for HTTP 429 (Rate Limit) or HTTP 403 (Access Denied) exception codes from the external APIs. If the primary fetch fails or times out, the fallback mechanism immediately intercepts the failure and pivots entirely to an internal synthetic data generator, ensuring zero frontend downtime.")
    doc.add_paragraph("Step 5 – Heuristic Calculation (AI Prediction Engine): Once the environmental data is successfully resolved, the system pipes the JSON structures into the algorithmic prediction modules. It calculates the Heat Index via the Rothfusz regression, computes dynamic hydration needs based on activity intensity, and evaluates physiological lung recovery time against PM2.5 concentrations.")
    doc.add_paragraph("Step 6 – Data Persistence & Validation: The output payload containing the health predictions is modeled using strict Mongoose schemas and safely executed as an asynchronous write operation on the MongoDB Atlas database cluster.")
    doc.add_paragraph("Step 7 – UI Rendering & State Updates: The final aggregated JSON payload is returned to the React frontend, where the React Virtual DOM is updated. The result is rendered to the user utilizing Framer Motion physics-based animations and Recharts SVGs for an intuitive, visual user experience.")

    doc.add_heading('3.2 Function Oriented Design', level=2)
    doc.add_paragraph("React Frontend Module: Implemented as a highly interactive SPA (Single Page Application), this module handles user input, renders glassmorphism-styled Tailwind components, and manages application routing via React Router. It acts as the gateway to all environmental tracking features.")
    doc.add_paragraph("Express Orchestration Engine: Residing in the Node.js backend (index.js), this module connects the REST routes to the underlying computational logic. It coordinates parameter passing, securely decrypts API keys from the .env file, and calls the appropriate sub-functions for environmental fetching.")
    doc.add_paragraph("Network API Gateway (Services Layer): This is the critical networking module (services.js) responsible for managing all Axios HTTP requests to external environmental providers. It implements the try-catch retry logic and autonomously manages the Mock Data Fallback mechanisms.")
    doc.add_paragraph("AI Prediction Module: Interfaces directly with the raw data to apply rigorous mathematical heuristics (predictions.js). It manages the strict isolation of complex formulas (like the Heat Index regression and Hydration logic) from the routing layer, ensuring pure, testable functions.")
    doc.add_paragraph("Data Persistence Module: Manages the reading, writing, and schema validation of MongoDB documents (PredictionLog) via Mongoose ODM (db.js). It ensures that all prediction states are safely persisted for historical auditing and provides graceful failure handling if the database connection drops.")

    doc.add_heading('3.3 System Design', level=2)
    doc.add_heading('3.3.1 Data Flow Diagrams', level=3)
    doc.add_paragraph("Level 0 – Context Diagram (Figure 3.1):")
    doc.add_paragraph("The Level 0 DFD represents the AtmoSense AI system as a central execution node interacting with the Host User, the External Weather Cloud (OpenWeatherMap/WAQI), and the MongoDB Database. The User provides physical context and location parameters; AtmoSense AI requests raw environmental data from the Cloud; it then calculates complex health heuristics, logs the transaction to the Database, and returns the actionable safety metrics back to the User's dashboard.")
    
    doc.add_paragraph("Level 1 – Architectural Workflow Diagram (Figure 3.2):")
    doc.add_paragraph("The Level 1 diagram breaks down the monolithic central node into discrete operational flows. The frontend React application passes state to the Express.js Router. The Router splits the execution thread: one thread initiates API calls to the WAQI service for PM2.5/PM10 data, while the other queries OpenWeatherMap for temperature and humidity. Upon resolution, both data streams converge into the Prediction Engine node. The computed results flow outward simultaneously to the MongoDB logging node and backward as an HTTP response to the Client UI node, where conditional logic renders RED/YELLOW/GREEN safety badges.")
    
    doc.add_paragraph("[Note: Insert visual Flowchart/DFD diagrams here matching the structural description above.]")

    output_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Detailed.docx'
    doc.save(output_path)
    print(f"Saved to {output_path}")

if __name__ == '__main__':
    append_chapter_3()
