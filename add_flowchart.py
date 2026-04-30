import base64
import urllib.request
import docx
from docx.shared import Inches
import os

def add_flowchart():
    # Mermaid diagram for AtmoSense Flow Chart (Fallback logic)
    mermaid_code = """
flowchart TD
    Start([Start API Fetch]) --> Dispatch[Dispatch Parallel Axios Requests]
    Dispatch --> CheckResp{HTTP Response Code}
    
    CheckResp -- 200 OK --> Parse[Parse JSON Payload]
    Parse --> Norm[Normalize Environmental Metrics]
    Norm --> Pred[Execute Heuristic Prediction Engine]
    
    CheckResp -- 403 or 429 --> Failover[Intercept Exception at Gateway]
    Failover --> Mock[Initialize Mock Data Generator]
    Mock --> Gen[Generate Bounded Synthetic Data]
    Gen --> Pred
    
    Pred --> Log[Log to MongoDB]
    Log --> End([Return Final JSON to UI])
    """
    
    b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}?type=png&bgColor=!white"
    
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0'}
    )
    
    img_path = r'Report\System_Flow_Chart.png'
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as out_file:
                out_file.write(response.read())
        print("Flow Chart image downloaded successfully.")
    except Exception as e:
        print(f"Failed to download image: {e}")
        return

    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_Ch2_Ch3_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.3.3 Flow Chart', level=3)
    
    doc.add_paragraph("The system Flow Chart provides a comprehensive, decision-based view of AtmoSense AI’s network resilience and execution logic, specifically focusing on the proprietary Mock Data Fallback architecture. Because the system relies on external environmental Application Programming Interfaces (APIs like OpenWeatherMap and WAQI) to process real-time calculations, it is inherently vulnerable to network instability and strict API usage quotas. The flow chart visually maps the intricate branching logic designed to mitigate these vulnerabilities and ensure zero-downtime execution for the end user.")
    doc.add_paragraph("The logical flow operates through a strict hierarchy of conditional decision nodes:")
    
    doc.add_paragraph("1. Initial Dispatch & Primary API Evaluation: The process flow begins when the Node.js core engine dispatches parallel Axios requests to the primary designated environmental APIs. The first major decision node evaluates the HTTP response code returned by the third-party servers.")
    doc.add_paragraph("2. Success Path (HTTP 200 OK): If the external servers return an HTTP 200 status, the flow chart branches directly to the normalization phase, where the system parses the JSON payload, validates the raw environmental metrics, and pipes the data into the Heuristic Prediction Engine for final analysis.")
    doc.add_paragraph("3. Exception Handling (HTTP 429 / HTTP 403): If the primary APIs reject the request due to quota exhaustion (HTTP 429 Rate Limit) or unauthorized access (HTTP 403 Forbidden), the flow chart redirects into the failover loop. Instead of crashing and returning a stack trace to the React frontend, the system intercepts the exception immediately at the network gateway layer.")
    doc.add_paragraph("4. Internal Mock Data Pivot: Upon exception interception, the flow chart depicts a hard pivot to a completely internal synthetic data generator. The system actively utilizes mathematical bounds to generate realistic meteorological and pollutant data (e.g., bounding temperature between local historical minimums and maximums) and injects this payload seamlessly into the execution pipeline. This loop ensures that the prediction heuristics always receive valid inputs, effectively masking any backend infrastructure failures or API outages from the user dashboard.")

    doc.add_paragraph('\nSystem Flow Chart Image:')
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph('\n')

    output_path = r'Report\AtmoSense_Ch1_to_Ch3_FullyCompleted.docx'
    doc.save(output_path)
    print("Flow chart text and image successfully appended to the document.")

if __name__ == '__main__':
    add_flowchart()
