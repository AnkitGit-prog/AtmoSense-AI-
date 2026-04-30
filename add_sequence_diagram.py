import base64
import urllib.request
import docx
from docx.shared import Inches
import os

def add_sequence_diagram():
    # Mermaid diagram for AtmoSense Sequence Flow
    mermaid_code = """
sequenceDiagram
    participant U as User
    participant F as React Frontend
    participant R as Express Router
    participant S as Services Layer
    participant API as External APIs
    participant DB as MongoDB

    U->>F: Clicks Predict Health Impact
    F->>F: Capture State
    F->>R: POST /api/predict
    
    R->>S: delegateFetch()
    S->>API: GET /weather and /aqi
    
    Note right of API: Rate Limit Scenario
    API-->>S: HTTP 429 (Rate Limit Exceeded)
    
    Note left of S: Intercept and Fallback
    S->>S: triggerMockFallback()
    S-->>R: Return Synthetic Data
    
    R->>R: computeHeuristics()
    R->>DB: savePredictionLog()
    DB-->>R: Confirmed Write
    
    R-->>F: HTTP 200 OK
    F-->>U: Render Safety Badges
    """
    
    b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}?type=png&bgColor=!white"
    
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0'}
    )
    
    img_path = r'Report\Sequence_Diagram.png'
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as out_file:
                out_file.write(response.read())
        print("Sequence Diagram image downloaded successfully.")
    except Exception as e:
        print(f"Failed to download image: {e}")
        return

    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch3_ERDiagram_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.3.6 Sequence Diagram', level=3)
    
    doc.add_paragraph("The Sequence Diagram illustrates the precise, time-ordered message exchange between the core components of the AtmoSense AI ecosystem. Unlike static structural diagrams, the sequence diagram maps the exact chronological lifecycle of a single prediction command—from human initiation to autonomous heuristic rendering—highlighting the critical role of the Backend Services Layer in mitigating API rate-limiting failures.")
    doc.add_paragraph("The diagram tracks the vertical timeline of execution across six interacting participants: the User, the React Frontend, the Express Backend (Router), the Services Layer, the External Environmental APIs (e.g., OpenWeatherMap), and the MongoDB Database.")
    doc.add_paragraph("The sequence unfolds in the following chronological phases:")
    
    doc.add_paragraph("1. Initiation & Context Capture: The sequence begins with a synchronous interaction from the User to the React Frontend, triggering the 'Predict Health Impact' mechanism. The Frontend executes a local state capture to gather the user's defined activity parameters and geolocation data.")
    doc.add_paragraph("2. Payload Transmission: The React application constructs the JSON payload and invokes an asynchronous POST request via Axios to the Express Backend API.")
    doc.add_paragraph("3. External API Failure (The Rate-Limit Scenario): The diagram models a worst-case network scenario. The Express Router delegates execution to the Services Layer, which dispatches a parallel GET request to the External APIs. In this sequence, the APIs respond with an HTTP 429 Rate Limit Error.")
    doc.add_paragraph("4. Autonomous Fallback Interception: Crucially, this error is not returned to the User. Instead, the Services Layer intercepts the failure internally. It immediately triggers its Mock Data Fallback mechanism, replacing the failed external data streams with bounded synthetic meteorological inputs.")
    doc.add_paragraph("5. Heuristic Computation & Validation: The Services Layer returns the synthetic payload back to the core router, which then passes the metrics through the pure functions of the Prediction Engine. The engine strictly computes the Heat Index, Hydration, and Lung Recovery values.")
    doc.add_paragraph("6. Persistence & Final Output: Once the heuristics are generated, the Express Router triggers an asynchronous write operation to the MongoDB Database via Mongoose. The sequence concludes as the final computed prediction object is passed back to the React Frontend as an HTTP 200 response, which then renders the color-coded safety badges to the User.")
    
    doc.add_paragraph("By visualizing the time-ordered messages, this Sequence Diagram mathematically proves the resilience of AtmoSense AI's architecture, demonstrating how complex multi-API failovers operate entirely in the background without interrupting the user's interactive workflow.")

    doc.add_paragraph('\nSystem Sequence Diagram Image:')
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph('\n')

    output_path = r'Report\AtmoSense_Ch1_to_Ch3_SequenceDiagram_Final.docx'
    doc.save(output_path)
    print("Sequence chart text and image successfully appended to the document.")

if __name__ == '__main__':
    add_sequence_diagram()
