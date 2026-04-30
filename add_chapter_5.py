import docx
import os

def add_chapter_5():
    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch4_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_page_break()

    # --- Chapter 5 ---
    doc.add_heading('Chapter 5\nResults and Discussions', level=1)
    
    doc.add_heading('5.1 User Interface Representation', level=2)
    
    doc.add_paragraph("The AtmoSense AI user interface is purposefully implemented as a highly interactive, Single Page Application (SPA) operating within the user's web browser. Rather than relying on a heavy, text-dense layout or a raw terminal output, the system leverages React.js, Tailwind CSS, and Framer Motion to provide a high-speed, visually intuitive, and distraction-free environmental tracking dashboard.")
    
    doc.add_paragraph("The main UI panel renders a modern glassmorphism-styled dashboard upon execution, distinctly featuring real-time weather metrics, an interactive health prediction form, and visual air quality gauges. The interface prompts users for inputs asynchronously (e.g., selecting physical activity intensity and health sensitivities from dropdown menus) and outputs real-time execution results.")
    
    doc.add_paragraph("Upon successful calculation via the backend API Prediction Engine, the frontend UI dynamically renders color-coded safety badges (e.g., Red for Danger, Yellow for Caution, Green for Safe), precise hydration requirements in ml/hr, and animated data charts directly onto the screen without requiring a page refresh. This ensures a seamless, modern user experience optimized for immediate cognitive recognition.")

    # Leaving space for images
    doc.add_paragraph("\n")
    p1 = doc.add_paragraph("[ PLACEHOLDER: Paste Screenshot of AtmoSense AI Main Dashboard Here ]")
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n\n\n")
    
    p2 = doc.add_paragraph("[ PLACEHOLDER: Paste Screenshot of Health Prediction Results Here ]")
    p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n\n\n")
    
    p3 = doc.add_paragraph("[ PLACEHOLDER: Paste Screenshot of Air Quality Analysis Here ]")
    p3.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n")

    output_path = r'Report\AtmoSense_Ch1_to_Ch5_Final.docx'
    doc.save(output_path)
    print("Chapter 5 successfully appended to the document.")

if __name__ == '__main__':
    add_chapter_5()
