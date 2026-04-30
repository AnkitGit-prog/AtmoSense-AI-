import base64
import urllib.request
import docx
from docx.shared import Inches
import os

def add_class_diagram():
    # Mermaid diagram for AtmoSense Class Architecture
    mermaid_code = """
classDiagram
    class AppFrontend {
        +Router Navigation
        +Global State
        +renderLayout()
    }
    class HomeView {
        -localState City
        +renderWeatherWidget()
        +fetchDashboardData()
    }
    class InputFormView {
        -activityType
        -healthSensitivity
        +handlePredictSubmit()
    }
    class ExpressServer {
        +port 5000
        +initRoutes()
        +handleRequests()
    }
    class ServicesLayer {
        -openWeatherKey
        -waqiKey
        +fetchWeather(city)
        +fetchAQI(city)
        +generateMockFallback()
    }
    class PredictionEngine {
        +calculateHeatIndex(temp, hum)
        +calculateHydration(activity)
        +estimateLungRecovery(aqi)
    }
    class PredictionLogSchema {
        +String city
        +Object weatherData
        +Object heuristics
        +saveToMongoDB()
    }

    AppFrontend --> HomeView
    AppFrontend --> InputFormView
    InputFormView --> ExpressServer
    HomeView --> ExpressServer
    ExpressServer --> ServicesLayer
    ExpressServer --> PredictionEngine
    PredictionEngine --> PredictionLogSchema
    """
    
    b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}?type=png&bgColor=!white"
    
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0'}
    )
    
    img_path = r'Report\Class_Diagram.png'
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as out_file:
                out_file.write(response.read())
        print("Class Diagram image downloaded successfully.")
    except Exception as e:
        print(f"Failed to download image: {e}")
        return

    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch3_FullyCompleted.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.3.4 Class Diagram', level=3)
    
    doc.add_paragraph("The Class Diagram provides a strict Object-Oriented and modular visual representation of the AtmoSense AI architecture. Because the application utilizes the MERN stack, the architecture is divided between frontend React component hierarchies and backend Node.js module encapsulations, ensuring maintainability, clean data abstraction, and strict separation of concerns between the user interface and the execution logic.")
    doc.add_paragraph("The diagram illustrates the core structural entities, their internal attributes, methods, and the exact dependency relationships that bind the system together:")
    
    doc.add_paragraph("1. The Frontend Orchestrator (App.jsx): This is the primary entry point and routing controller of the client application. It handles UI state management, navigation routing, and the main presentation layout. The App component acts as a parent, possessing a dependency relationship with the underlying page components.")
    doc.add_paragraph("2. The View Components (Home, InputForm, AirQuality): To ensure modularity, the system relies on distinct view classes (React functional components). Each view encapsulates its local state and specialized methods—for example, the Home component encapsulates methods for renderWeatherWidget(), while the InputForm encapsulates handlePredictSubmit() and input validation.")
    doc.add_paragraph("3. The Backend Orchestration Engine (Express Server): The index.js module acts as the core server class. It handles incoming HTTP traffic and routes it to specific functional controllers. It guarantees secure API routing, allowing the system to execute any endpoint request interchangeably.")
    doc.add_paragraph("4. The Network Gateway (Services Layer): This critical utility module operates independently of the core router but maintains a strict association relationship with it. The router instantiates the Services layer to handle external HTTP communications. It encapsulates private attributes for API keys (openWeatherKey, waqiKey) and exposes public methods for data fetching and fallback generation, ensuring that the route controllers remain entirely agnostic to network routing complexities.")
    doc.add_paragraph("5. The Prediction Engine: This module encapsulates pure heuristic algorithms. It exposes public mathematical methods like calculateHeatIndex() and calculateHydration(), taking raw environmental state as arguments and returning deterministic objects without side effects.")
    doc.add_paragraph("6. Data Structures (PredictionLog Schema): To manage complex data persistence efficiently, the system utilizes Mongoose Schemas. The PredictionLog schema formally type-hints and encapsulates meteorological data, user profiles, and generated predictions to ensure data integrity within MongoDB before persistence.")
    
    doc.add_paragraph("By utilizing this strict modular design, the Class Diagram demonstrates how AtmoSense AI maintains a decoupled, scalable architecture where new environmental tracking features can be seamlessly integrated simply by extending the existing routing and service layers.")

    doc.add_paragraph('\nSystem Class Diagram Image:')
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph('\n')

    output_path = r'Report\AtmoSense_Ch1_to_Ch3_ClassDiagram_Final.docx'
    doc.save(output_path)
    print("Class chart text and image successfully appended to the document.")

if __name__ == '__main__':
    add_class_diagram()
