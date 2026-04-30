import base64
import urllib.request
import docx
from docx.shared import Inches
import os

def add_er_diagram():
    # Mermaid diagram for AtmoSense ER Diagram
    mermaid_code = """
erDiagram
    PREDICTION_LOG {
        ObjectId id PK
        Date timestamp
    }
    USER_PROFILE {
        String activity_type
        String health_sensitivity
        String location_city
    }
    METEOROLOGICAL_DATA {
        Float temperature_c
        Float humidity_percent
        Float pm25_concentration
        Float pm10_concentration
    }
    HEALTH_HEURISTICS {
        Float heat_index_c
        Float hydration_ml_hr
        String lung_recovery_time
        String safety_level
    }

    PREDICTION_LOG ||--|| USER_PROFILE : encapsulates
    PREDICTION_LOG ||--|| METEOROLOGICAL_DATA : encapsulates
    PREDICTION_LOG ||--|| HEALTH_HEURISTICS : encapsulates
    """
    
    b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}?type=png&bgColor=!white"
    
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0'}
    )
    
    img_path = r'Report\ER_Diagram.png'
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as out_file:
                out_file.write(response.read())
        print("ER Diagram image downloaded successfully.")
    except Exception as e:
        print(f"Failed to download image: {e}")
        return

    # Open existing document
    doc_path = r'Report\AtmoSense_Ch1_to_Ch3_ClassDiagram_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('3.3.5 ER Diagram', level=3)
    
    doc.add_paragraph("The Entity-Relationship (ER) Diagram illustrates the logical data schema and cardinality of the AtmoSense AI ecosystem. Unlike conventional enterprise applications that rely on heavy relational SQL databases with strict normal forms, AtmoSense AI is explicitly designed for high-speed, flexible NoSQL execution using MongoDB and Mongoose ODM. In this context, the \"entities\" in the ER diagram represent structured BSON document collections stored persistently in the Atlas Cloud database.")
    doc.add_paragraph("The diagram maps the structural relationships between the core data entities required for the prediction and tracking modules:")
    
    doc.add_paragraph("1. The Core Entity (USER_PROFILE): The central entity of the system during a prediction event is the user's physiological state. It exhibits a one-to-one (1:1) relationship with the generated health predictions. It stores user-defined attributes such as activity_type (e.g., Running, Gym) and health_sensitivity (e.g., Asthma, Normal).")
    doc.add_paragraph("2. Environmental Context (METEOROLOGICAL_DATA): This entity is strictly bound to the prediction event. It stores the aggregated real-time weather and pollution metrics fetched from external APIs, such as temperature, humidity, PM2.5, and PM10 values. This strict isolation ensures that the raw environmental inputs are independently verifiable.")
    doc.add_paragraph("3. AI Computations (HEALTH_HEURISTICS): Associated with the Prediction Engine, this complex sub-entity stores the calculated physiological outputs resulting from the raw data. It contains specific health metrics like heat_index_celsius, hydration_ml_per_hour, and lung_recovery_hours. It maintains a 1:1 relationship with the USER_PROFILE, mapping the exact environmental impact to the specific user's physical condition.")
    doc.add_paragraph("4. The Aggregate Record (PREDICTION_LOG): To optimize historical auditing and data persistence, the system utilizes the PREDICTION_LOG entity as the master wrapper collection. This entity encapsulates the USER_PROFILE, METEOROLOGICAL_DATA, and HEALTH_HEURISTICS within a single document hierarchy. It contains metadata attributes like timestamp and primary _id.")
    
    doc.add_paragraph("By visualizing these entities, the ER Diagram demonstrates AtmoSense AI's highly efficient approach to data persistence. The structured Mongoose schemas guarantee strict data integrity for the frontend rendering engines, while utilizing MongoDB's nested document architecture eliminates the latency and overhead associated with managing traditional relational database JOIN operations.")

    doc.add_paragraph('\nSystem ER Diagram Image:')
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph('\n')

    output_path = r'Report\AtmoSense_Ch1_to_Ch3_ERDiagram_Final.docx'
    doc.save(output_path)
    print("ER chart text and image successfully appended to the document.")

if __name__ == '__main__':
    add_er_diagram()
