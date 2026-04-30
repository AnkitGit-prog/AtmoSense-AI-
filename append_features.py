import docx
import os

def append_features_snapshots():
    doc_path = r'Report\AtmoSense_Ch1_to_Ch5_Final.docx'
    if not os.path.exists(doc_path):
        print("Error: Could not find document.")
        return

    doc = docx.Document(doc_path)
    
    doc.add_heading('5.2 Brief Description of Various Modules of the System', level=2)
    
    doc.add_paragraph("React View Modules (Dashboard / AQI Monitor / Predictor): These core modules interact natively with the browser DOM via React components and Tailwind CSS. They monitor user inputs, visualize real-time temperature, and render complex Framer Motion physics-based animations to display dynamic safety gauges.", style='List Bullet')
    
    doc.add_paragraph("Backend Orchestration Module (Express.js Router): This module intercepts all incoming HTTP requests from the React frontend. It securely handles Cross-Origin Resource Sharing (CORS), decrypts API keys from the local environment, and delegates parallel fetches to external providers.", style='List Bullet')
    
    doc.add_paragraph("API Gateway & Fallback Module (Services Layer): A highly resilient networking component that dispatches requests to OpenWeatherMap and WAQI. It acts as a safety net by intercepting HTTP 429 Rate Limit and 403 Forbidden exceptions, instantly routing the data flow to an internal Mock Data Generator to prevent frontend crashes.", style='List Bullet')
    
    doc.add_paragraph("AI Heuristic Engine (Prediction Algorithms): A collection of pure mathematical functions that compute complex physiological impacts. It utilizes the Rothfusz regression for calculating the Heat Index and scales hydration requirements based on the user's selected physical activity (e.g., Running vs. Walking).", style='List Bullet')
    
    doc.add_paragraph("Data Persistence Module (MongoDB Atlas): A NoSQL database integration that validates incoming prediction payloads against strict Mongoose schemas. It records the historical context of each prediction event, including the user's city and safety level, ensuring auditable historical logs.", style='List Bullet')
    
    doc.add_heading('5.3 System Operations: Snapshots', level=2)
    doc.add_paragraph("The following section illustrates the precise algorithmic flow and execution snapshot for each of the core features of the AtmoSense AI system:")

    features = [
        ("Feature 1: Real-Time Meteorological Dashboard", 
         "Description: Upon initial application load, this feature fetches the current geolocation parameters and queries the OpenWeatherMap API. The React frontend dynamically renders the ambient temperature, wind speed, and humidity in a highly visual, glassmorphism-styled widget."),
        
        ("Feature 2: Air Quality & Pollutant Analyzer", 
         "Description: This module queries the WAQI (World Air Quality Index) API to extract granular pollutant concentrations (PM2.5, PM10, NO2, O3). It visually maps these metrics onto interactive graphical gauges, clearly indicating the severity of outdoor air pollution to the user."),
        
        ("Feature 3: AI Health Impact Prediction", 
         "Description: The core predictive feature of the application. The user inputs their desired physical activity (e.g., Gym) and personal health sensitivity (e.g., Asthma). The system merges these inputs with the real-time weather and AQI data, mathematically calculates the Heat Index and Lung Recovery time, and renders explicit safety badges (Safe/Caution/Danger)."),
        
        ("Feature 4: Mock Data Fallback Architecture", 
         "Description: A backend resilience feature. If the user disconnects from the internet or the OpenWeatherMap API exhausts its rate limits, this module intercepts the HTTP exception and generates realistic, bounded synthetic data. This ensures the frontend dashboard continues to function flawlessly during API outages."),
        
        ("Feature 5: Database Persistence Logging", 
         "Description: Whenever a prediction is successfully generated, the Express backend triggers an asynchronous write operation to the MongoDB Atlas cluster. The system logs the exact timestamp, environmental context, and physiological output, ensuring a permanent historical record.")
    ]

    for title, desc in features:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
        p = doc.add_paragraph(f"\n[ PLACEHOLDER: Paste Screenshot of {title.split(': ')[1]} Here ]\n")
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n\n\n\n")

    output_path = r'Report\AtmoSense_Ch1_to_Ch5_Complete_Final.docx'
    doc.save(output_path)
    print("Features and Snapshots successfully appended to Chapter 5.")

if __name__ == '__main__':
    append_features_snapshots()
